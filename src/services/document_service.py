"""Document generation service for Word templates."""

import logging
import re
from pathlib import Path
from typing import Dict, Optional, Any
from datetime import datetime

from docx import Document
from docx.shared import Pt

from config.settings import RESUME_TEMPLATE_DRIVE_ID
from .drive_service import DriveService

PLACEHOLDER_PATTERN = re.compile(r'\{\{[^}]+\}\}')

logger = logging.getLogger(__name__)


class DocumentService:
    """Handle Word template processing."""

    def __init__(self, drive_file_id: Optional[str] = None):
        """Initialize document service.

        Args:
            drive_file_id: Google Drive file ID for template

        Raises:
            FileNotFoundError: If template cannot be found or downloaded
        """
        self.drive_file_id = drive_file_id or RESUME_TEMPLATE_DRIVE_ID
        if not self.drive_file_id:
            raise ValueError("RESUME_TEMPLATE_DRIVE_ID must be configured")
        
        # Initialize drive service
        self.drive_service = DriveService()
        logger.info(f"Document service initialized with Google Drive template: {self.drive_file_id}")

    
    def _download_template_from_drive(self) -> str:
        """Download template from Google Drive.
        
        Returns:
            Path to downloaded template file
            
        Raises:
            Exception: If download fails
        """
        try:
            # Create temp directory for downloaded template
            temp_dir = Path("./temp")
            temp_dir.mkdir(exist_ok=True)
            temp_file = temp_dir / "resume_template.docx"
            
            # Download template
            logger.info(f"Downloading template from Google Drive: {self.drive_file_id}")
            downloaded_path = self.drive_service.download_file_binary(self.drive_file_id, str(temp_file))
            
            return downloaded_path
            
        except Exception as error:
            logger.error(f"Failed to download template: {error}")
            if "404" in str(error):
                raise FileNotFoundError(
                    f"Resume template not found (ID: {self.drive_file_id}). "
                    "Please ensure:\n"
                    "1. The template file exists in Google Drive\n"
                    "2. The file ID in your .env file is correct\n"
                    "3. The file is shared with your service account or made public"
                )
            elif "403" in str(error):
                raise PermissionError(
                    f"Access denied to template (ID: {self.drive_file_id}). "
                    "Please ensure the file is shared with your service account email."
                )
            else:
                raise Exception(f"Failed to download template: {error}")
    

    def generate_resume(
        self,
        resume_sections: Dict[str, Any],
        job_metadata: Dict,
        output_dir: str = "./generated_resumes"
    ) -> str:
        """Generate resume document from template.

        Args:
            resume_sections: Dict containing role bullets and skills
            job_metadata: Job metadata dict
            output_dir: Directory to save generated files

        Returns:
            Path to generated .docx file

        Raises:
            Exception: If document generation fails
        """
        try:
            # Create output directory
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)

            # Download and load template
            template_path = self._download_template_from_drive()
            doc = Document(template_path)

            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company = job_metadata.get("company", "Unknown").replace(" ", "_")
            title = job_metadata.get("title", "Position").replace(" ", "_")
            
            # Sanitize filename to remove problematic characters
            company = re.sub(r'[/\\:*?"<>|]', '_', company)
            title = re.sub(r'[/\\:*?"<>|]', '_', title)
            
            filename = f"Resume_{company}_{title}_{timestamp}.docx"
            output_file = output_path / filename

            # Replace placeholders
            self._replace_placeholders(doc, resume_sections, job_metadata)

            # Save document
            doc.save(str(output_file))
            logger.info(f"Generated resume document: {output_file}")

            return str(output_file)

        except Exception as error:
            logger.error(f"Error generating resume document: {error}")
            raise

    def _replace_text_in_paragraph(self, paragraph, placeholder: str, replacement_text: str) -> None:
        """Replace text in a paragraph while preserving run structure.

        Args:
            paragraph: python-docx Paragraph object
            placeholder: Text to find and replace
            replacement_text: Text to replace with
        """
        # Check if placeholder exists in paragraph text
        if placeholder not in paragraph.text:
            return

        while placeholder in paragraph.text:
            start_index = paragraph.text.index(placeholder)
            end_index = start_index + len(placeholder)

            current_pos = 0
            start_run_idx = None
            start_offset = None
            end_run_idx = None
            end_offset = None

            for idx, run in enumerate(paragraph.runs):
                run_text = run.text
                run_len = len(run_text)

                if start_run_idx is None and current_pos + run_len > start_index:
                    start_run_idx = idx
                    start_offset = start_index - current_pos

                if end_run_idx is None and current_pos + run_len >= end_index:
                    end_run_idx = idx
                    end_offset = end_index - current_pos
                    break

                current_pos += run_len

            if start_run_idx is None or end_run_idx is None:
                logger.warning("Placeholder boundaries could not be resolved in paragraph; aborting replacement")
                return

            start_run = paragraph.runs[start_run_idx]
            end_run = paragraph.runs[end_run_idx]

            prefix = start_run.text[:start_offset] if start_offset is not None else start_run.text
            suffix = end_run.text[end_offset:] if end_offset is not None else ""

            start_run.text = f"{prefix}{replacement_text}{suffix}"

            for idx in range(start_run_idx + 1, end_run_idx + 1):
                paragraph.runs[idx].text = ""
    
    def _remove_placeholders_from_paragraph(self, paragraph) -> None:
        """Strip any unreplaced placeholders without disturbing paragraph style."""
        for run in paragraph.runs:
            if PLACEHOLDER_PATTERN.search(run.text):
                run.text = PLACEHOLDER_PATTERN.sub('', run.text)
    
    def _apply_bullet_style(self, paragraph) -> None:
        """Apply a bullet style to the given paragraph if available in the template."""
        try:
            paragraph.style = 'List Bullet'
            logger.info("Applied List Bullet style")
        except KeyError:
            try:
                paragraph.style = 'ListBullet'
                logger.info("Applied ListBullet style")
            except KeyError:
                logger.warning("Bullet style not found in template; falling back to inline bullet character")
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    if first_run.text and not first_run.text.lstrip().startswith(("•", "-", "*")):
                        first_run.text = f"• {first_run.text}"

    def _replace_placeholders(
        self,
        doc: Document,
        resume_sections: Dict[str, Any],
        job_metadata: Dict
    ) -> None:
        """Replace placeholders in Word document.

        Args:
            doc: python-docx Document object
            resume_sections: Dict containing role bullets and skills
            job_metadata: Job metadata dict
        """
        # Build replacement map
        replacements = {}
        bullet_placeholders = set()  # Track which placeholders are bullets

        # Add role-specific bullet replacements
        for role, content in resume_sections.items():
            # Skip the skills entry as it's not a role
            if role == "skills" or not isinstance(content, list):
                continue

            role_bullets = content
            logger.info(f"Processing {len(role_bullets)} bullets for role: {role}")

            for i, bullet in enumerate(role_bullets):
                # Clean the bullet text first (remove any leading bullet characters)
                clean_bullet = bullet.lstrip('•*- ').strip()
                placeholder = f"{{{{{role}_EXPERIENCE_BULLET_{i+1}}}}}"
                # Store just the clean bullet text - Word will add the bullet marker
                replacements[placeholder] = clean_bullet
                bullet_placeholders.add(placeholder)

            # Clear any remaining placeholders for this role
            for i in range(len(role_bullets), 10):  # Assume max 10 bullets per role
                placeholder = f"{{{{{role}_EXPERIENCE_BULLET_{i+1}}}}}"
                replacements[placeholder] = ""

        # Clear placeholders for any missing roles
        expected_roles = ["LEAFICIENT", "DHS", "EDUCATIVE_PM", "EDUCATIVE_SWE"]
        for role in expected_roles:
            if role not in resume_sections or not isinstance(resume_sections.get(role), list):
                logger.warning(f"No bullets found for role: {role}, clearing placeholders")
                for i in range(10):
                    placeholder = f"{{{{{role}_EXPERIENCE_BULLET_{i+1}}}}}"
                    replacements[placeholder] = ""

        # Add other replacements
        # Use skills extracted by the LLM from the candidate's resume
        skills_text = resume_sections.get("skills", "Python, Machine Learning, Data Analysis")

        replacements.update({
            "{{CANDIDATE_NAME}}": "Your Name",  # TODO: Add to config
            "{{CONTACT_INFO}}": "your.email@example.com | (123) 456-7890 | LinkedIn",
            "{{SUMMARY}}": self._generate_summary(job_metadata),
            "{{TECHNICAL_SKILLS}}": skills_text,
        })

        # Log replacements for debugging
        logger.info(f"Replacing placeholders: {len(replacements)} total")
        for key in sorted(replacements.keys()):
            if "BULLET" in key and replacements[key]:
                logger.info(f"  {key}: {replacements[key][:50]}...")

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    logger.info(f"Found placeholder {key} in paragraph")
                    self._replace_text_in_paragraph(paragraph, key, value)

                    if key in bullet_placeholders and value:
                        self._apply_bullet_style(paragraph)

        # Remove any remaining placeholders while preserving styles
        for paragraph in doc.paragraphs:
            self._remove_placeholders_from_paragraph(paragraph)

        # Replace in tables (cell paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                logger.info(f"Found placeholder {key} in table cell")
                                self._replace_text_in_paragraph(paragraph, key, value)

                                if key in bullet_placeholders and value:
                                    self._apply_bullet_style(paragraph)
        
        # Remove any remaining placeholders from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._remove_placeholders_from_paragraph(paragraph)

    def _generate_summary(self, job_metadata: Dict) -> str:
        """Generate professional summary based on job.

        Args:
            job_metadata: Job metadata dict

        Returns:
            Summary text
        """
        # Simple template-based summary
        role_level = job_metadata.get("role_level", "experienced")
        title = job_metadata.get("title", "professional")

        return (
            f"{role_level} professional with proven expertise in "
            f"{title.lower()} and related technologies. "
            "Strong track record of delivering results in fast-paced environments."
        )


    def cleanup_files(self, *file_paths: str) -> None:
        """Delete temporary files.

        Args:
            *file_paths: Variable number of file paths to delete
        """
        for file_path in file_paths:
            try:
                path = Path(file_path)
                if path.exists():
                    path.unlink()
                    logger.info(f"Deleted temporary file: {file_path}")
            except Exception as error:
                logger.warning(f"Failed to delete {file_path}: {error}")
