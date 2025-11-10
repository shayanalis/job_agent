"""LLM service for OpenAI and Claude interactions."""

import logging
from typing import List, Dict, Optional, Any

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import JsonOutputParser

from config.settings import (
    OPENAI_API_KEY,
    OPENAI_MODEL,
    LLM_TEMPERATURE,
)
from src.agents.state import AnalyzedRequirements, ValidationResult, JobMetadata
from typing import Tuple
from src.prompts.prompt_templates import (
    JD_ANALYZER_SYSTEM_PROMPT,
    JD_ANALYZER_USER_PROMPT,
    METADATA_EXTRACTOR_SYSTEM_PROMPT,
    METADATA_EXTRACTOR_USER_PROMPT,
    COMPLETE_RESUME_VALIDATOR_SYSTEM_PROMPT,
    format_complete_resume_validator_prompt,
    RESUME_REWRITER_SYSTEM_PROMPT,
    RESUME_REWRITER_USER_PROMPT,
)

logger = logging.getLogger(__name__)


class LLMService:
    """Wrapper for LLM API calls with structured outputs."""

    def __init__(self):
        """Initialize LLM clients."""
        if not OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY not found in environment variables")

        # Certain models only support the default temperature of 1.0
        default_temp_models = (
            "gpt-5",
            "gpt-5-preview",
            "gpt-4.1",
            "gpt-4.1-mini",
        )

        requires_default_temp = any(
            OPENAI_MODEL.startswith(prefix) for prefix in default_temp_models
        )

        kwargs = {
            "model": OPENAI_MODEL,
            "api_key": OPENAI_API_KEY,
        }

        if requires_default_temp:
            if LLM_TEMPERATURE not in (None, 1, 1.0):
                logger.warning(
                    "Model %s requires default temperature=1.0; overriding configured value %s",
                    OPENAI_MODEL,
                    LLM_TEMPERATURE,
                )
            kwargs["temperature"] = 1.0
        elif LLM_TEMPERATURE is not None:
            kwargs["temperature"] = LLM_TEMPERATURE

        self.openai_model = ChatOpenAI(**kwargs)

        logger.info(
            "LLM Service initialized with %s (temperature=%s)",
            OPENAI_MODEL,
            kwargs.get("temperature", "default"),
        )

    def analyze_job_description(self, job_description: str, job_metadata: Dict) -> AnalyzedRequirements:
        """Analyze job description and extract structured requirements.

        Args:
            job_description: Raw job description text
            job_metadata: Job metadata dict

        Returns:
            AnalyzedRequirements object with extracted data

        Raises:
            Exception: If LLM call fails
        """
        system_prompt = JD_ANALYZER_SYSTEM_PROMPT.format()
        user_prompt = JD_ANALYZER_USER_PROMPT.format(
            job_title=job_metadata.get('title', 'Not specified'),
            company=job_metadata.get('company', 'Not specified'),
            job_description=job_description
        )

        try:
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]

            response = self.openai_model.invoke(messages)
            parser = JsonOutputParser(pydantic_object=AnalyzedRequirements)
            result = parser.parse(response.content)

            # Validate with Pydantic
            analyzed = AnalyzedRequirements(**result)
            logger.info(f"Successfully analyzed JD: {len(analyzed.required_skills)} required skills found")
            return analyzed

        except Exception as error:
            logger.error(f"Error analyzing job description: {error}")
            raise

    def extract_job_metadata(self, job_description: str, provided_metadata: Dict) -> JobMetadata:
        """Extract or enhance job metadata from description.

        Args:
            job_description: Raw job description text
            provided_metadata: Metadata provided by Chrome extension

        Returns:
            JobMetadata object

        Raises:
            Exception: If LLM call fails
        """
        system_prompt = METADATA_EXTRACTOR_SYSTEM_PROMPT.format()
        user_prompt = METADATA_EXTRACTOR_USER_PROMPT.format(
            job_description=job_description,
            provided_title=provided_metadata.get('title', 'Unknown'),
            provided_company=provided_metadata.get('company', 'Unknown'),
            provided_url=provided_metadata.get('job_url', '')
        )

        try:
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]

            response = self.openai_model.invoke(messages)
            parser = JsonOutputParser(pydantic_object=JobMetadata)
            result = parser.parse(response.content)

            metadata = JobMetadata(**result)
            logger.info(f"Extracted metadata: {metadata.title} at {metadata.company}")
            return metadata

        except Exception as error:
            logger.error(f"Error extracting metadata: {error}")
            raise ValueError(f"Failed to extract job metadata from description: {str(error)}")

    def validate_complete_resume(
        self,
        docx_file_path: str,
        job_description: Optional[str] = None,
        job_metadata: Optional[Dict] = None,
        requirements: Optional[AnalyzedRequirements] = None
    ) -> ValidationResult:
        """Validate a complete resume document from DOCX file.

        Args:
            docx_file_path: Path to the DOCX resume file
            job_description: Optional original job description
            job_metadata: Optional job metadata dict
            requirements: Optional analyzed requirements

        Returns:
            ValidationResult with comprehensive feedback

        Raises:
            Exception: If file reading or LLM call fails
        """
        from docx import Document
        from pathlib import Path
        
        # Read DOCX file
        try:
            doc = Document(docx_file_path)
            
            # Extract text from document
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            resume_content = "\n".join(full_text)
            
        except Exception as error:
            logger.error(f"Error reading DOCX file: {error}")
            raise ValueError(f"Failed to read resume file: {str(error)}")
        
        # Build validation prompt based on what's provided
        system_prompt = COMPLETE_RESUME_VALIDATOR_SYSTEM_PROMPT.format()
        
        # Convert requirements to dict if it's an object
        requirements_dict = None
        if requirements:
            requirements_dict = {
                'required_skills': requirements.required_skills,
                'keywords_for_ats': requirements.keywords_for_ats
            }
        
        user_prompt = format_complete_resume_validator_prompt(
            resume_content=resume_content,
            job_description=job_description,
            job_metadata=job_metadata if isinstance(job_metadata, dict) else (job_metadata.model_dump() if job_metadata else None),
            requirements=requirements_dict
        )

        try:
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]

            response = self.openai_model.invoke(messages)
            parser = JsonOutputParser(pydantic_object=ValidationResult)
            result = parser.parse(response.content)

            validation = ValidationResult(**result)
            logger.info(f"Complete resume validation: {validation.is_valid}, score: {validation.keyword_coverage_score}")
            return validation

        except Exception as error:
            logger.error(f"Error validating complete resume: {error}")
            raise

    def analyze_job_complete(
        self, 
        job_description: str, 
        provided_metadata: Dict
    ) -> Tuple[JobMetadata, AnalyzedRequirements]:
        """Analyze job description and extract both metadata and requirements in one call.

        Args:
            job_description: Raw job description text
            provided_metadata: Metadata provided by Chrome extension

        Returns:
            Tuple of (JobMetadata, AnalyzedRequirements)

        Raises:
            Exception: If LLM call fails
        """
        # Create a combined prompt that extracts both metadata and requirements
        system_prompt = """You are an expert job description analyzer. Extract both job metadata 
        and requirements in a single analysis. Be comprehensive and precise."""
        
        user_prompt = f"""Analyze this job description and extract all information:

Job Description:
{job_description}

Provided Metadata:
- Title: {provided_metadata.get('title', 'Not specified')}
- Company: {provided_metadata.get('company', 'Not specified')}
- URL: {provided_metadata.get('url', 'Not specified')}

Return a JSON object with two sections:
{{
    "metadata": {{
        "title": "extracted or use provided",
        "company": "extracted or use provided",
        "role_level": "Entry/Mid/Senior/Staff/Principal/Lead/Manager/Director/VP/C-Level/Not Specified",
        "sponsorship": "Yes/No/Not Specified - Look for phrases like 'must be authorized to work', 'no sponsorship', 'unable to sponsor', 'H1B', 'visa sponsorship', 'work authorization required'",
        "posted_date_raw": "e.g., 'Posted 3 days ago' or ''",
        "job_type": "Full-time/Part-time/Contract/Internship or null",
        "job_url": "{provided_metadata.get('url', '')}"
    }},
    "requirements": {{
        "required_skills": ["skill1", "skill2"],
        "preferred_skills": ["skill1", "skill2"],
        "soft_skills": ["skill1", "skill2"],
        "key_responsibilities": ["responsibility1", "responsibility2"],
        "must_have_experience": ["experience1", "experience2"],
        "nice_to_have": ["item1", "item2"],
        "domain_knowledge": ["domain1", "domain2"],
        "years_experience_required": 5,
        "education_requirements": "Bachelor's degree or equivalent",
        "certifications": ["cert1", "cert2"],
        "keywords_for_ats": ["keyword1", "keyword2"]
    }}
}}"""

        try:
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]

            response = self.openai_model.invoke(messages)
            parser = JsonOutputParser()
            result = parser.parse(response.content)
            
            # Parse metadata and requirements
            metadata_dict = result.get('metadata', {})
            requirements_dict = result.get('requirements', {})
            
            # Create validated objects
            metadata = JobMetadata(**metadata_dict)
            requirements = AnalyzedRequirements(**requirements_dict)
            
            logger.info(f"Successfully analyzed job: {metadata.title} at {metadata.company}")
            logger.info(f"Found {len(requirements.required_skills)} required skills, {len(requirements.keywords_for_ats)} ATS keywords")
            
            return metadata, requirements

        except Exception as error:
            logger.error(f"Error in combined job analysis: {error}")
            raise ValueError(f"Failed to analyze job description: {str(error)}")

    def rewrite_resume(
        self,
        base_resume_pointers: str,
        requirements: AnalyzedRequirements,
        validation_feedback: Optional[str] = None
    ) -> Dict[str, List[str]]:
        """Transform base resume pointers into job-specific bullet points.

        Args:
            base_resume_pointers: Raw experience descriptions from pointer file
            requirements: Analyzed job requirements
            validation_feedback: Optional feedback from validation to improve transformation

        Returns:
            Dict mapping role names to lists of transformed bullet points

        Raises:
            Exception: If LLM call fails
        """
        logger.info("Starting rewrite_resume method")
        logger.debug(f"base_resume_pointers length: {len(base_resume_pointers) if base_resume_pointers else 0}")
        
        try:
            system_prompt = RESUME_REWRITER_SYSTEM_PROMPT.format()
            
            # Prepare validation feedback section
            validation_feedback_section = ""
            if validation_feedback:
                validation_feedback_section = f"Previous Validation Feedback: {validation_feedback}"

            # Convert requirements to a simple string format
            analyzed_requirements_str = requirements.model_dump_json(indent=2)
            
            user_prompt = RESUME_REWRITER_USER_PROMPT.format(
                analyzed_requirements=analyzed_requirements_str,
                base_resume_pointers=base_resume_pointers,
                validation_feedback_section=validation_feedback_section
            )
        except Exception as e:
            logger.error(f"Error preparing prompts: {e}")
            raise

        try:
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]

            response = self.openai_model.invoke(messages)
            
            # Log the raw response for debugging
            logger.debug(f"Raw LLM response: {response.content[:500]}...")
            
            # First try to parse as JSON
            try:
                parser = JsonOutputParser()
                result = parser.parse(response.content)
            except Exception as parse_error:
                logger.error(f"Failed to parse JSON response: {parse_error}")
                logger.error(f"Raw response: {response.content}")
                
                # Try to extract JSON from the response if it's wrapped in text
                import json
                
                # Try to find the start of JSON object in the response
                content = response.content.strip()
                
                # If the response starts with some text before the JSON, find the opening brace
                start_idx = content.find('{')
                if start_idx != -1:
                    # Try to parse from the first opening brace
                    json_str = content[start_idx:]
                    
                    # Find the matching closing brace
                    brace_count = 0
                    end_idx = -1
                    for i, char in enumerate(json_str):
                        if char == '{':
                            brace_count += 1
                        elif char == '}':
                            brace_count -= 1
                            if brace_count == 0:
                                end_idx = i + 1
                                break
                    
                    if end_idx != -1:
                        try:
                            json_to_parse = json_str[:end_idx]
                            result = json.loads(json_to_parse)
                            logger.info("Successfully extracted JSON from response text")
                        except json.JSONDecodeError as e:
                            logger.error(f"JSON decode error: {e}")
                            raise ValueError(f"Could not parse JSON from response: {json_to_parse[:200]}...")
                    else:
                        raise ValueError("Could not find matching closing brace for JSON")
                else:
                    raise ValueError(f"No JSON object found in response: {content[:200]}...")
            
            # The response should be a dictionary with role names as keys
            if not isinstance(result, dict):
                raise ValueError(f"Response is not a dictionary, got {type(result)}")
            
            # Extract skills if present
            skills = result.pop("skills", None)
            if skills:
                logger.info(f"Extracted skills: {skills[:100]}...")
            
            # Validate the structure for role-specific bullets
            role_specific_bullets = {}
            for role, bullets in result.items():
                if not isinstance(bullets, list):
                    raise ValueError(f"Bullets for role {role} is not a list")
                # Ensure all items are strings
                role_specific_bullets[role] = [str(bullet) for bullet in bullets]
            
            # Add skills to the result if present
            if skills:
                role_specific_bullets["skills"] = skills
            
            total_bullets = sum(len(bullets) for role, bullets in role_specific_bullets.items() if role != "skills")
            logger.info(f"Generated {total_bullets} tailored bullet points across {len([r for r in role_specific_bullets if r != 'skills'])} roles")
            
            return role_specific_bullets

        except Exception as error:
            logger.error(f"Error rewriting resume: {error}")
            raise

